#!/usr/bin/python
#coding:utf-8
# tested in Win
# Version: 20190109


import time,os,sys
from docx import Document
from docx.shared import Inches
import comtypes.client

def sir():
    wp = os.path.dirname(os.path.realpath(__file__))
    # os.chdir(wp)
    name = input('Server Name >> ')
    SN = input('Serial Number >> ')
    site = input('Press Enter for default Shanghai >> ')
    if site == '': site = 'Shanghai'
    loc = input('eg:IDX2/A24/GSHFRB024A2/Bay7 >>> ')
    SIR = os.path.join(wp,'SIR Blade.docx')
    out_file = os.path.join(wp,'%s.pdf' % name)

    """
    Serial Number:          6CU830G700
    Site:   SHANGHAI
    Area/Building/Room:
    IDX2/A24/GSHFRB024A2/Bay9
    """

    doc = Document(SIR)
    tables = doc.tables

    t0 = tables[0]
    cName = t0.cell(0,0)
    cName.text = 'System Name:      '+name
    cSN = t0.cell(0,1)
    cSN.text = 'Serial Number:          '+SN
    cSite = t0.cell(1,0)
    cSite.text = 'Site:   '+site
    cloc = t0.cell(1,1)
    cloc.text = 'Area/Building/Room:\n'+loc

    t1 = tables[3]
    cdate = t1.cell(1,1)
    cdate.text = time.strftime('%Y-%m-%d',time.localtime(time.time()))

    # sig = r'E:\UT\signature.png'
    # s = t1.cell(1,2).add_paragraph()
    # r = s.add_run()
    # r.add_picture(sig,width = Inches(0.5), height = Inches(0.5) )

    try:
        doc.save(SIR)
    except PermissionError as e:
        print(e)
        print('Is file being opened?')

    # os.startfile(SIR,'print')

    wdFormatPDF = 17
    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(SIR)
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()

if __name__=='__main__':
    try:
        sir()
    except KeyboardInterrupt:
        print('ctrl + c')

